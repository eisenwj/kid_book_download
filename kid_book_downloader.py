# coding=GBK
import sys
import requests
import os
import io
import re
import bs4
import lxml
import PIL
import docx
import ctypes
import multiprocessing
import concurrent.futures
from bs4 import BeautifulSoup
from PIL import Image
from asyncio import streams
from docx import Document
from docx.enum.section import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches
from docx.image.exceptions import *


#get the kidbook's title
def gettitle(soup):
    
    results=soup.find_all('meta')

    for each in results:
        property_value=each.get("property")
        if (property_value == "og:title"):
            content_value=each.get("content")
            title = str(content_value).strip(" <>.:?|*\"/").replace("|","").replace(":","").replace("?","").replace("/", "")
            return title
            break
    return("None")

#get all the pictures' urls' list
def img_search(soup):
    #results=soup.find_all('img')
    results=soup.select('img')
    
    return results

#check and process picture downloaded and make it all horizoned if rotate_mode is 90/270 degree
#Or keep it portraited if rotate_mode is 0 degree
def img_validate(pathname,rotate_mode):

    img=Image.open(pathname)

    #convert the png or other format with gray scale to jpeg
    if img.mode in ("RGBA", "P"): 
        img = img.convert("RGB")

    width=img.size[0]
    height=img.size[1]

    if (height > width) and (rotate_mode > 0):
        img = img.transpose(rotate_mode)
        img.save(pathname)
        img.close()
        return float(height/width)
    return float(width/height)

#set the page orientation to be either portrait or landscape
#orientation should be "portraited"/"landscape"
def page_set_orientation(doc,run,orientation):
    
    current_section = doc.sections[-1]
    current_width, current_height = current_section.page_height, current_section.page_width
    
    if (orientation == "landscape"):
        if (current_width < current_height):
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            new_section.orientation = WD_ORIENTATION.PORTRAIT
            new_section.page_width = current_height
            new_section.page_height = current_width
            p = doc.add_paragraph()
            p.alignment= WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            
    if (orientation == "portrait"):
        if (current_width > current_height):
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            new_section.orientation = WD_ORIENTATION.LANDSCAPE
            new_section.page_width = current_height
            new_section.page_height = current_width
            p = doc.add_paragraph()
            p.alignment= WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
    
    return run
    

#download one picture from assigned url
def pic_download(dir,pic_name,pic_url):
    try:
        fp=open(dir+os.sep+pic_name, 'wb')
        fp.write(requests.get(pic_url).content)
        fp.close()
#        img_validate(dir+os.sep+pic_name)
        return 0
    except requests.exceptions.ConnectionError:
        return -1

#return one kid book's soup object
def page_parser(url):
    #    r = requests.get("https://mp.weixin.qq.com/s?__biz=MzA5OTE2NjUyMw==&mid=2666327871&idx=3&sn=0ddf7205a8dd82d7456d5a5a31c7c8b7&chksm=8b9e3f6bbce9b67d03bff9d854a0d708d09f7bf9dc1036550be1ca1889c812381e08a57ea864&token=262424641&lang=zh_CN&scene=21#wechat_redirect")
    r = requests.get(url)
    
    soup = BeautifulSoup(r.text,'lxml')
    return soup


#check if the picture is an advertise or not
#if is advertise -- return true, else return false
def img_filter(url_str):
    adv_list=[
        r"http://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aAW3OcBzKXH6yibzoMgKrSMZOS99xQaD0uwRdiaTTH8ibSApDhuTZ6WBMqJgM8ZtMJ6qiaYjCtzWxlU7Q/0?wxfrom=5&wx_fmt=gif",
        r"https://mmbiz.qpic.cn/mmbiz_gif/xvjlzfjg30HfgVpxOmBXv3Wib4qDXFH8sanWcl7605GBD09PcSwIial58iaQ4fqzdvXHAbIe2YfsF9RAKWrO3eISw/640?wx_fmt=gif",
        r"https://mmbiz.qpic.cn/mmbiz_png/wHtT7l3B4aCrjUWduXMEPZEwRQucZpfjvxvyOXmJXVYP5nDKqicW69Y396k4Ps0uVq8HmHZjB1D592dmuoJDHWQ/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/wHtT7l3B4aADm9QHjicSv7ejN0lHNJUKuQAMXIK0xXxdSyMGvVxB5fYx2VK40z8icvOW6rh9mIlichibpQzndZzlRQ/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/wHtT7l3B4aCrjUWduXMEPZEwRQucZpfjpfwJicdR0rq5vE9pnw2zOHAvTLlyic929v6nEvykXAg2GiaicVStsn8sjQ/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_png/wHtT7l3B4aCrjUWduXMEPZEwRQucZpfjvxvyOXmJXVYP5nDKqicW69Y396k4Ps0uVq8HmHZjB1D592dmuoJDHWQ/640?wx_fmt=png",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/jR58jqMGnUZSdAHRAcmVWBbXzVDpicpvZ8PYvtneibicybticMicuE9ia70MDUkTVA4Bgf9NoPDIL70icPibuVz6ibnO94A/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_png/xvjlzfjg30FZJf153icaVD901QsA9LMAjJh0bqHyfY7EAKIiaP2DicynIFHTarN4QsiaHYYmmiaQyZ0XuC1t4JW6gPg/640?wx_fmt=png",
        r"https://mmbiz.qpic.cn/mmbiz_png/xvjlzfjg30FicI1vxgZQBAU7OicZORSrCEo0WqhXvib714QwEW7g2bcQliapA2abhXqRgyQqV0WGFmhFRg3bwg8Cog/640?wx_fmt=png",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/6bML8pV2ozEcG6JV3zgb6ibOJsxic0ic8vib12RCDwW957m2WciazBUNiblUt8cuefUo2aa6XibGz7xZSM52ZnE9r4FAQ/640?wx_fmt=jpeg",
        r"http://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aDymGIgvnib6KouFAAk8HSDfxHNfQDwlLPhr02WVyOF7ZYdxXjESBibeqDtEdWo1THcyOuT1xF76CWA/640?wx_fmt=jpeg",
        r"http://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aAvz6sHVz7uq5I8BAjH23QI7wp8PF9FP7kL3x5iaLgHPJhtgSqc5q9MAOAsyGjicRA7478O2ia0gicItg/640?wx_fmt=png",
        r"http://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aC88sEefGxFhofjuh9Nr2AhibkGJuK0T1kaSnefVibduibCy1u80qrg0xdTV60p5aGQUbsBib7GOCticHw/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aDymGIgvnib6KouFAAk8HSDfxHNfQDwlLPhr02WVyOF7ZYdxXjESBibeqDtEdWo1THcyOuT1xF76CWA/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz/wHtT7l3B4aC88sEefGxFhofjuh9Nr2AhibkGJuK0T1kaSnefVibduibCy1u80qrg0xdTV60p5aGQUbsBib7GOCticHw/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/wHtT7l3B4aBaas6nn0kZWiaARODl83I2UQ6ianx0duIbmr3s2uppyb52Z2Bos2RVuILljC2SHSSo5vneS6xmb5Ug/640?wx_fmt=jpeg",
        r"https://mmbiz.qpic.cn/mmbiz_jpg/6bML8pV2ozF8thsia5m9FvGn0vVcibfJ1rxF1YBagyz0YGrr06VQLvyqYrFnVdV2J91ITVCGicK8quCGVg88ojr2Q/640?wx_fmt=jpeg"
        ]
    if (url_str in adv_list):
        return True
    else:
        return False


#download all pictures from a single book and use them to build up a docx file 
def kidbook_download(kidbook_url,errormsgs,configures):

    soup=page_parser(kidbook_url)
    
    #get the book's title to create folder with book name
    book_title=gettitle(soup)
    if (book_title == "None"):
        errormsgs.append("小人书名未检索到，请检查代码或url: "+kidbook_url)
        print("小人书名未检索到，请检查代码或url:  "+kidbook_url+"\n")
        return -10
    
    dir=str(configures[0])
    rotate_mode=int(configures[1])
    img_dir=dir+os.sep+book_title
    
    if not os.path.isdir(img_dir):
        os.mkdir(img_dir)
    
    if (str(sys.platform) == "win32"):
        doc = Document(r"C:\Temp\Print\小人书\01_模板.docx")
    else:
        if (str(sys.platform) == "linux"):
            doc = Document(r"/mnt/c/Temp/Print/小人书/01_模板.docx")
    
    #check if need portrait page orientation
    if ( rotate_mode == -1 ):
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
#        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        current_section.orientation = WD_ORIENT.PORTRAIT
        current_section.page_width = new_width
        current_section.page_height = new_height

#        current_section = doc.sections[-1]
#        new_width, new_height = current_section.page_height, current_section.page_width
#        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
#        new_section.orientation = WD_ORIENTATION.PORTRAIT
#        new_section.page_width = new_width
#        new_section.page_height = new_height

    
    
    p = doc.add_paragraph()
    p.alignment= WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(book_title+" url: "+kidbook_url)
    p = doc.add_paragraph()
    p.alignment= WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()

    imgs=img_search(soup)
    

    i=0
    for each_img in imgs:
        pic_url=str(imgs[i].get("data-src"))
        print(str(i)+":   "+book_title+"    "+pic_url)
        if ((pic_url == "None") 
#            or (pic_url.find(r"640?wx_fmt=")==-1)
            or (img_filter(pic_url))
#            or (pic_url.find(r"wx_fmt=png")!=-1)
#            or (pic_url.find(r"640?tp=webp")!=-1)
#            or (pic_url.find(r"wx_fmt=gif")!=-1)
            ):
            i=i+1
            continue


        for j in range(5):
#           flag=0 if download successfully, and !0 if failed 
            flag = pic_download(img_dir,"640_"+str(i)+".jpg",pic_url)

            if (flag == 0):
                wh_ratio=float(img_validate(img_dir+os.sep+"640_"+str(i)+".jpg",rotate_mode))
                try:
                #A4 paper size is 297*210
                    if (rotate_mode > 0 ):
                        if wh_ratio > float(297.0/210.0):
                            run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",width=Inches(11))
                        else:
                            run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",height=Inches(8))
                    else:
                        if (wh_ratio >= 1):
                            #a picture in landscape in a portraited book, need to rotate the current page only
                            run=page_set_orientation(doc, run, "landscape")
                            if wh_ratio > float(297.0/210.0):
                                run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",width=Inches(11))
                            else:
                                run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",height=Inches(8))
                            run=page_set_orientation(doc, run, "portrait")
                        else:
                            if wh_ratio < float(210.0/297.0):
                                run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",height=Inches(11))
                            else:
                                run.add_picture(img_dir+os.sep+"640_"+str(i)+".jpg",width=Inches(8))
                    break
                except docx.image.exceptions.UnexpectedEndOfFileError:
                    print("640_"+str(i)+".jpg failed to add to doc.")
                    errormsgs.append("640_"+str(i)+".jpg failed to add to doc.")
                    break
            else:
                print("640_"+str(i)+".jpg failed for "+str(j)+" times.")
                if (j == 4):
                    errormsgs.append("This book got picture lost: "+kidbook_url)
                    errormsgs.append("640_"+str(i)+".jpg for "+book_title+" failed:"+pic_url)
                    print("5 retries failed. Connection is lost for: "+pic_url)
        
        i=i+1
    
    doc.save(dir+os.sep+book_title+".docx")
    return 0

def book_download_process(tasks:multiprocessing.Queue,errormsgs,configures):
    pid=os.getpid();
    print("Sub-Process "+str(pid)+" started")
#    print("Tasks:"+str(args.qsize()))
    while not (tasks.empty()):
        kidbook=tasks.get_nowait()
        print("Process "+str(pid)+" : book--"+kidbook)
#        time.sleep(1+ random.random()*5)#simulate the downloading time
        kidbook_download(str(kidbook),errormsgs,configures)
        
    return

    
    
#download a pack of single books--usually "合集" -- eg. "说岳全传"合集
def kidbook_pack_download(url,errormsgs,configures):
    tasks=mgr.Queue()
    
    r = requests.get(url)
    soup = BeautifulSoup(r.text,'lxml')
    results=soup.find_all('a')
    i=0
    for a in results:
        if (a.get("data-linktype") in ["1","2"]):
            kidbook_url=a.get("href")
            print(str(i)+":  "+kidbook_url)
            tasks.put_nowait(kidbook_url)
#            kidbook_download(kidbook_url)
            i=i+1
    
    with concurrent.futures.ProcessPoolExecutor(max_workers=3) as executor:
        f=[executor.submit(book_download_process,tasks,errormsgs,configures) for i in range(3)]
        concurrent.futures.wait(f) 
    
    return 0


if (__name__  ==  "__main__") :
    #get the book's url--
    #Next action will added book pack url
    
    #contains all the error messages to print at last
    mgr = multiprocessing.Manager()
    errormsgs=mgr.list()
    configures=mgr.list()
    #Folder for book downloading
#    dir=mgr.Value(ctypes.c_char_p,"")
    dir=""

    print("Started:")
    url=input("请输入小人书 网页链接:")
    is_pack=int(input("是否是合集--1:合集；2:单本    "))

    while is_pack not in [1,2]:
        is_pack=int(input("是否是合集--1:合集；2:单本    "))

    if (sys.platform=="win32"):
        dir=input(r"请输入下载目录:[C:\Temp\Movie\downthemall]  ")
    else:
        if (sys.platform=="linux"):
            dir=input(r"请输入下载目录:[/mnt/c/Temp/Movie/downthemall]  ")
        
    if (dir == ""):
        if (sys.platform=="win32"):
            dir=r"C:\Temp\Movie\downthemall"
        else:
            if (sys.platform=="linux"):
                dir=r"/mnt/c/Temp/Movie/downthemall"

    rotate_direction=input(r"如需旋转，旋转方向: 1. 逆时针; 2. 顺时针;3.不旋转:   [1]")
    while rotate_direction not in ['','1','2','3']:
        rotate_direction=input(r"如需旋转，旋转方向: 1. 逆时针; 2. 顺时针:  [1]")
    
    if ((rotate_direction == '1') or (rotate_direction=="")) :
        rotate_mode=Image.Transpose.ROTATE_90
    if (rotate_direction == '2') :
        rotate_mode=Image.Transpose.ROTATE_270
    if (rotate_direction == '3') :
        rotate_mode=-1

    configures.append(dir)
    configures.append(rotate_mode)

    if (is_pack == 2):
        print("单本下载")
        kidbook_download(url,errormsgs,configures)
    else:
        print("合集下载")
        kidbook_pack_download(url,errormsgs,configures)

    
    for msg in errormsgs:
        print(msg)
